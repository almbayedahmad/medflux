from typing import Optional
from docx import Document

def read_docx_to_text(path: str) -> str:
    doc = Document(path)
    parts = []
    # Paragraphs
    for p in doc.paragraphs:
        txt = (p.text or "").rstrip()
        if txt:
            parts.append(txt)
    # Tables (optional: row by row)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [ (c.text or "").strip() for c in row.cells ]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


