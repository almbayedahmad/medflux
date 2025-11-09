# PURPOSE:
#   DOCX parsing helpers for readers stage.
#
# OUTCOME:
#   Extracts normalized text from DOCX files for downstream readers logic.
#
# INPUTS:
#   - Path to a DOCX document.
#
# OUTPUTS:
#   - Extracted text as a single normalized string (in-memory).
#
# DEPENDENCIES:
#   - Optional: python-docx (Document) for reading DOCX content.
from __future__ import annotations

"""DOCX parsing helpers for the readers runtime."""

from typing import List

try:
    from docx import Document  # type: ignore
except Exception as exc:  # pragma: no cover - optional dependency
    Document = None  # type: ignore
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


def get_docx_text(path: str) -> str:
    """Return normalized text extracted from a DOCX document."""

    if Document is None:  # pragma: no cover - dependency missing at runtime
        raise RuntimeError("python-docx is required to read DOCX files") from DOCX_IMPORT_ERROR

    doc = Document(path)
    parts: List[str] = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").rstrip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def get_readers_docx_text(path: str) -> str:
    """Alias retained for compatibility with legacy imports."""

    return get_docx_text(path)


__all__ = ["get_docx_text", "get_readers_docx_text"]
